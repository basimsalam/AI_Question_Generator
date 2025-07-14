from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
import os
from datetime import datetime
from fpdf import FPDF

from pathlib import Path


def save_question_paper_as_word(paper: dict, filename="question_paper.docx") -> str:
    doc = Document()
    doc.add_heading("Generated Question Paper", 0)

    for key, question in paper.items():
        doc.add_paragraph(f"{key}: {question}", style="List Number")

    filepath = os.path.join("generated_papers", filename)
    os.makedirs("generated_papers", exist_ok=True)
    doc.save(filepath)
    return filepath

def save_question_paper_as_pdf(paper: dict, filename="question_paper.pdf") -> str:
    output_dir = Path("generated_papers")
    output_dir.mkdir(exist_ok=True)
    filepath = output_dir / filename

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    pdf.set_title("Generated Question Paper")
    pdf.set_font("Arial", "B", 14)
    pdf.cell(0, 10, "Generated Question Paper", ln=True, align="C")
    pdf.ln(10)

    pdf.set_font("Arial", "", 12)
    for q_no, question in paper.items():
        wrapped_lines = pdf.multi_cell(0, 10, f"{q_no}: {question}", align="L")
        pdf.ln(4)  # space between questions

    pdf.output(str(filepath))
    return str(filepath)
